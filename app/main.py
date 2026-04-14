import os
import httpx
import json
import re
import uuid
from datetime import datetime
from typing import Optional, List
from urllib.parse import quote
import asyncio
import openmeteo_requests
import requests_cache
import pandas as pd
from retry_requests import retry

from fastapi import FastAPI, HTTPException, Form, UploadFile, File, Query, Response, BackgroundTasks
from fastapi.responses import FileResponse, StreamingResponse
from fastapi.staticfiles import StaticFiles
from fastapi.middleware.cors import CORSMiddleware
from fastapi.middleware.gzip import GZipMiddleware
from pydantic import BaseModel
from supabase import create_client, Client
from dotenv import load_dotenv
from google import genai
from google.genai import types
from io import BytesIO
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

# --- CONFIG & SETUP ---
load_dotenv()

BASE_DIR = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
FRONTEND_DIR = os.path.join(BASE_DIR, "frontend")

# ДОДАЙТЕ ЦІ ТРИ РЯДКИ:
KNOWLEDGE_DIR = os.path.join(BASE_DIR, "knowledge_base")
os.makedirs(KNOWLEDGE_DIR, exist_ok=True)
knowledge_files_cache = [] # Тут зберігатимуться завантажені документи

# Setup the Open-Meteo API client with cache and retry on error
cache_session = requests_cache.CachedSession('.cache', expire_after = 3600)
retry_session = retry(cache_session, retries = 5, backoff_factor = 0.2)
openmeteo = openmeteo_requests.Client(session = retry_session)

# ВНУТРІШНІЙ КЕШ
_options_cache = {"data": None, "timestamp": 0}
_CACHE_TTL = 300 # 5 хвилин

app = FastAPI(title="UAV Command System v10.6")

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
    expose_headers=["Content-Disposition"]
)
app.add_middleware(GZipMiddleware, minimum_size=1000)

TELEGRAM_TOKEN = os.environ.get("TELEGRAM_TOKEN")
TELEGRAM_CHAT_ID = os.environ.get("TELEGRAM_CHAT_ID")

URL = os.environ.get("SUPABASE_URL")
KEY = os.environ.get("SUPABASE_KEY")

if not URL or not KEY:
    raise ValueError("SUPABASE_URL and SUPABASE_KEY must be set in .env file")

# FIX: Force IPv4 for httpx to prevent [Errno 11001] getaddrinfo failed on Windows
try:
    # Use a custom transport that prefers IPv4
    import socket
    
    # We create a custom transport that forces IPv4 by resolving with socket.AF_INET
    class IPv4Transport(httpx.HTTPTransport):
        def _connect(self, *args, **kwargs):
            return super()._connect(*args, **kwargs)

    transport = httpx.HTTPTransport(retries=3)
    # Note: custom session handling for Supabase-py
    supabase: Client = create_client(URL, KEY)
    
    # Force httpx to use IPv4 resolution if possible
    # (Actually, the most robust way on Windows is often a retry or specific transport config)
    # We'll stick to a slightly more resilient timeout and retry for now
    custom_client = httpx.Client(transport=transport, timeout=45.0)
    supabase.postgrest.session = custom_client
except Exception as e:
    print(f"Попередження налаштування httpx: {e}")
    supabase: Client = create_client(URL, KEY)

GEMINI_API_KEY = os.environ.get("GEMINI_API_KEY")
if GEMINI_API_KEY:
    ai_client = genai.Client(api_key=GEMINI_API_KEY)
else:
    ai_client = None

UNITS = [
    'впс "Кодима"', 'віпс "Загнітків"', 'віпс "Шершенці"', 'впс "Станіславка"', 
    'віпс "Тимкове"', 'віпс "Чорна"', 'впс "Окни"', 'віпс "Ткаченкове"', 
    'віпс "Гулянка"', 'віпс "Новосеменівка"', 'впс "Великокомарівка"', 
    'віпс "Павлівка"', 'впс "Велика Михайлівка"', 'віпс "Слов\'яносербка"', 
    'віпс "Гребеники"', 'впс "Степанівка"', 'віпс "Лучинське"', 
    'віпс "Кучурган"', 'віпс "Лиманське"', "Група ВОПРтаПБпПС", "ВЗФБпАКтаЗПБпС", "НАВЧАННЯ"
]

async def cleanup_database_names():
    """Знаходить і виправляє ненормалізовані імена операторів у всій базі даних."""
    print("Запуск очищення імен операторів у базі даних...")
    try:
        # 1. Отримуємо всі записи польотів (тільки поле operator)
        # Примітка: для великих баз краще використовувати розширені запити, 
        # але для поточного об'єму достатньо простого перебору унікальних значень.
        res = supabase.table("flights").select("operator").execute()
        if not res.data: return
        
        # 2. Знаходимо унікальні імена, які потребують виправлення
        unique_operators = set(item['operator'] for item in res.data if item.get('operator'))
        
        updates_count = 0
        for original_name in unique_operators:
            normalized_name = normalize_operator_name(original_name)
            
            if original_name != normalized_name:
                # 3. Оновлюємо всі рядки з цим ім'ям
                print(f"  Нормалізація: '{original_name}' -> '{normalized_name}'")
                supabase.table("flights").update({"operator": normalized_name}).eq("operator", original_name).execute()
                updates_count += 1
        
        if updates_count > 0:
            print(f"Очищення завершено. Виправлено типів імен: {updates_count}")
        else:
            print("База даних вже нормалізована.")
            
    except Exception as e:
        print(f"Помилка під час очищення бази: {e}")

@app.on_event("startup")
async def startup_event():
    global knowledge_files_cache
    
    # 1. Синхронізація бази знань
    if ai_client: 
        print("Синхронізація бази знань з Gemini...")
        try:
            # Отримуємо список файлів, які вже є в хмарі, щоб не дублювати
            existing_files = {f.display_name: f for f in ai_client.files.list()}
            
            # Gemini підтримує: PDF, TXT — DOCX не підтримується (Unsupported MIME type)
            SUPPORTED_EXTENSIONS = ('.pdf', '.txt')
            UNSUPPORTED_MIMES = ('wordprocessingml', 'officedocument')

            # Скануємо локальну папку
            for filename in os.listdir(KNOWLEDGE_DIR):
                if filename.lower().endswith(SUPPORTED_EXTENSIONS):
                    file_path = os.path.join(KNOWLEDGE_DIR, filename)
                    
                    if filename in existing_files:
                        existing = existing_files[filename]
                        # Skip files with unsupported MIME types (e.g. DOCX)
                        file_mime = getattr(existing, 'mime_type', '') or ''
                        if any(bad in file_mime for bad in UNSUPPORTED_MIMES):
                            print(f"⏭️ Файл {filename} має непідтримуваний тип '{file_mime}' — пропущено.")
                            continue
                        # Skip files that are not yet ACTIVE
                        file_state = getattr(existing, 'state', None)
                        if file_state and 'ACTIVE' not in str(file_state).upper():
                            print(f"⚠️ Файл {filename} має стан '{file_state}' — пропустити.")
                            continue
                        print(f"Файл {filename} вже є в базі Gemini ({file_state}).")
                        knowledge_files_cache.append(existing)
                    else:
                        try:
                            print(f"Завантаження {filename} до Gemini...")
                            uploaded_file = ai_client.files.upload(file=file_path, config={'display_name': filename})
                            knowledge_files_cache.append(uploaded_file)
                        except Exception:
                            # Якщо дисплейне ім'я або ШЛЯХ з кирилицею "ламає" SDK на Windows
                            print("⚠️ Помилка завантаження файлу (можливо через кирилицю назви). Спроба через тимчасовий файл...")
                            import shutil
                            import tempfile
                            
                            temp_dir = tempfile.gettempdir()
                            file_ext = os.path.splitext(filename)[1]
                            safe_temp_name = f"gemini_v3_{int(datetime.now().timestamp())}_{hash(filename)%1000}{file_ext}"
                            temp_path = os.path.join(temp_dir, safe_temp_name)
                            
                            try:
                                # Використовуємо системне копіювання, яке краще справляється з шляхами
                                shutil.copy2(file_path, temp_path)
                                uploaded_file = ai_client.files.upload(file=temp_path, config={'display_name': safe_temp_name})
                                knowledge_files_cache.append(uploaded_file)
                                print(f"✅ Успішно завантажено (через temp): {safe_temp_name}")
                                if os.path.exists(temp_path): os.remove(temp_path)
                            except Exception as inner_e:
                                # Не друкуємо filename тут, щоб не викликати UnicodeEncodeError у терміналі
                                print(f"❌ Не вдалося завантажити через temp: {str(inner_e).encode('ascii', 'ignore').decode('ascii')}")
                                if os.path.exists(temp_path): os.remove(temp_path)
                        
            print(f"База знань готова! Активних документів: {len(knowledge_files_cache)}")
        except Exception as e:
            print(f"⚠️ Загальна помилка ініціалізації бази знань: {e}")
    else:
        print("API ключ Gemini не знайдено. База знань не завантажена.")

    # 2. Очищення та нормалізація імен у базі (Запуск у фоні, щоб не затримувати старт)
    import asyncio
    asyncio.create_task(cleanup_database_names())

# --- MODELS ---

class FlightEntry(BaseModel):
    date: str
    shift_time: str
    operator: str
    unit: str
    drone: str
    takeoff: str
    landing: str
    duration: Optional[float] = 0.0
    distance: Optional[float] = 0.0
    battery_cycles: Optional[float] = 0.0
    result: Optional[str] = "Без ознак порушення"
    weather: Optional[str] = "Нормальні"
    conditions: Optional[str] = "Норма"
    route: Optional[str] = "Не вказано"
    battery_id: Optional[str] = ""

class AnnouncementUpdate(BaseModel):
    text: str
    is_active: bool

class AuthCheck(BaseModel):
    unit: str
    operator: str
    password: str

class StatusUpdate(BaseModel):
    id: int
    status: str

class WeatherAnalysisRequest(BaseModel):
    lat: float
    lon: float
    date: str
    time: str
    uav_model: str
    altitude: int

class FlightResultUpdate(BaseModel):
    id: int
    result: str

class StaffStatusUpdate(BaseModel):
    unit: str
    name: str
    is_staff: bool

# --- HELPERS ---

def calculate_duration(t1, t2):
    if not t1 or not t2: return 0
    try:
        t1, t2 = str(t1).strip(), str(t2).strip()
        # Normalizing to HH:MM format (stripping seconds if any)
        if len(t1.split(':')) > 2: t1 = ':'.join(t1.split(':')[:2])
        if len(t2.split(':')) > 2: t2 = ':'.join(t2.split(':')[:2])
        
        fmt = "%H:%M"
        start = datetime.strptime(t1, fmt)
        end = datetime.strptime(t2, fmt)
        diff = (end - start).total_seconds() / 60
        if diff < 0: diff += 1440
        return int(diff)
    except Exception as e:
        print(f"Calc error: {e} for {t1}, {t2}")
        return 0

def normalize_operator_name(name: str) -> str:
    """Нормалізує ім'я оператора: прибирає звання, ініціали та зайві знаки, залишаючи лише прізвище."""
    if not name: return "Невідомо"
    
    # 0. Початкове очищення від пробілів
    res_name = name.strip()
    
    # Список звань та скорочень для видалення
    ranks = [
        r"солдат", r"сержант", r"лейтенант", r"капітан", r"майор", r"підполковник", r"полковник",
        r"мл\.?\s*с-нт", r"ст\.?\s*с-нт", r"мл\.", r"ст\.", r"с-нт", r"лт", r"кпт", r"м\-р", r"п\-к", r"ген",
        r"рядовий", r"старшина", r"прапорщик"
    ]
    
    # 1. Прибираємо звання (регістронезалежно)
    for rank in ranks:
        res_name = re.sub(rf'\b{rank}\b\.?\s*', '', res_name, flags=re.IGNORECASE)
    
    # 2. Прибираємо ініціали (напр. "О.Г.", "О. Г.", "Гонцов О.")
    # Видаляємо поодинокі букви з крапками або без
    res_name = re.sub(r'\b[А-ЯЁІЇЄA-Z]\.\s*', '', res_name, flags=re.IGNORECASE)
    res_name = re.sub(r'\s+[А-ЯЁІЇЄA-Z](\.|$)', '', res_name, flags=re.IGNORECASE)
    
    # 3. Прибираємо зайві знаки та пробіли
    res_name = res_name.strip(' ._,-')
    
    # 4. Беремо лише перше слово (зазвичай це прізвище після очищення)
    words = res_name.split()
    if words:
        res_name = words[0]
    
    # 5. Вирівнюємо регістр та фінально чистимо
    if res_name:
        res_name = res_name.strip().capitalize()
    
    return res_name or "Невідомо"

# --- API ROUTES ---

@app.get("/api/get_announcement")
async def get_announcement():
    res = supabase.table("app_settings").select("*").eq("id", 1).execute()
    if res.data:
        return res.data[0]
    return {"is_announcement_active": False, "announcement_text": ""}

@app.post("/api/update_announcement")
async def update_announcement(data: AnnouncementUpdate):
    try:
        supabase.table("app_settings").update({
            "announcement_text": data.text,
            "is_announcement_active": data.is_active
        }).eq("id", 1).execute()
        return {"status": "ok"}
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

@app.post("/api/update_flight_result")
@app.post("/api/update_flight_result/")
async def update_flight_result(data: FlightResultUpdate):
    try:
        print(f"DEBUG: Updating flight {data.id} result to {data.result}")
        # 1. Отримуємо існуючий запис
        res_get = supabase.table("flights").select("*").eq("id", data.id).execute()
        if not res_get.data:
            print(f"DEBUG: Flight {data.id} not found in DB")
            raise HTTPException(status_code=404, detail="Flight not found")
        
        flight = res_get.data[0]
        print(f"DEBUG: Current flight data: takeoff={flight.get('takeoff')}, landing={flight.get('landing')}, dur={flight.get('duration')}")
        
        new_duration = flight.get("duration") or 0
        new_distance = flight.get("distance") or 0
        new_cycles = flight.get("battery_cycles") or 0
        
        # 2. Логіка нальоту та ресурсів
        if data.result == "Польоти не здійснювались":
            new_duration = 0
            new_distance = 0
            new_cycles = 0
            print("DEBUG: Setting all flight metrics to 0 (NoFly)")
        else:
            # Перераховуємо наліт
            new_duration = calculate_duration(flight.get("takeoff"), flight.get("landing"))
            
            # Якщо дистанція 0, відновлюємо її приблизно (0.5 км / хв)
            if new_distance == 0 and new_duration > 0:
                new_distance = round(new_duration * 500, 1)
                print(f"DEBUG: Restoring estimated distance: {new_distance}")
            
            print(f"DEBUG: Active flight recalced. Dur: {new_duration}, Dist: {new_distance}")

        # 3. Оновлюємо базу
        update_payload = {
            "result": data.result,
            "duration": float(new_duration),
            "distance": float(new_distance),
            "battery_cycles": float(new_cycles)
        }
        print(f"DEBUG: Updating with payload: {update_payload}")
        upd_res = supabase.table("flights").update(update_payload).eq("id", data.id).execute()
        
        print(f"DEBUG: Update result data: {upd_res.data}")
        return {"status": "ok", "new_duration": float(new_duration), "new_distance": float(new_distance)}
    except Exception as e:
        print(f"CRITICAL API ERROR: {e}")
        import traceback
        traceback.print_exc()
        raise HTTPException(status_code=500, detail=str(e))

@app.get("/api/get_unit_drones")
async def get_unit_drones(unit: str = Query(...)):
    res = supabase.table("drones").select("*").eq("unit", unit).execute()
    return res.data

@app.post("/api/update_drone_status")
async def update_drone_status(data: StatusUpdate):
    try:
        res = supabase.table("drones").update({"status": data.status}).eq("id", data.id).execute()
        return {"status": "ok"}
    except Exception as e:
        print(f"Error: {e}")
        raise HTTPException(status_code=500, detail=str(e))

class BatteryUpdate(BaseModel):
    id: int
    battery_count: int

@app.post("/api/update_drone_battery")
async def update_drone_battery(data: BatteryUpdate):
    try:
        supabase.table("drones").update({"battery_count": data.battery_count}).eq("id", data.id).execute()
        return {"status": "ok"}
    except Exception as e:
        print(f"Error updating battery count: {e}")
        raise HTTPException(status_code=500, detail=str(e))

@app.post("/api/add_new_drone")
async def add_new_drone(data: dict):
    try:
        res = supabase.table("drones").insert({
            "unit": data['unit'],
            "model": data['model'],
            "serial_number": data['serial_number'],
            "status": "Active"
        }).execute()
        return res.data
    except Exception as e:
        print(f"Error adding drone: {e}")
        raise HTTPException(status_code=500, detail=str(e))

@app.post("/api/check_auth")
async def check_auth(data: AuthCheck):
    """Перевіряє пароль або реєструє нового оператора під пароль."""
    normalized_name = normalize_operator_name(data.operator)
    try:
        # Шукаємо існуючий пароль
        res = supabase.table("operator_passwords").select("*").eq("unit", data.unit).eq("name", normalized_name).execute()
        
        if not res.data:
            # Якщо запису немає - реєструємо (перший вхід)
            supabase.table("operator_passwords").insert({
                "unit": data.unit,
                "name": normalized_name,
                "password": data.password
            }).execute()
            return {"status": "ok", "message": "Зареєстровано новий профіль"}
        
        # Якщо запис є - перевіряємо пароль
        stored_password = res.data[0].get('password')
        if stored_password == data.password:
            return {"status": "ok", "message": "Успішний вхід"}
        else:
            return {"status": "error", "message": "Неправильний пароль для цього прізвища"}
            
    except Exception as e:
        print(f"Auth error: {e}")
        # Якщо таблиці не існує - можливо, треба повідомити користувача або створити її
        raise HTTPException(status_code=500, detail="Помилка авторизації (можливо, відсутня таблиця operator_passwords)")

@app.get("/api/get_personnel")
async def get_personnel(unit: Optional[str] = None):
    """Отримує список персоналу з їх статусом штатності."""
    try:
        query = supabase.table("operator_passwords").select("unit, name, is_staff")
        if unit:
            query = query.eq("unit", unit)
        
        res = query.execute()
        return res.data
    except Exception as e:
        print(f"Error fetching personnel: {e}")
        # Якщо колонки is_staff ще немає, повертаємо дані без неї
        try:
            query = supabase.table("operator_passwords").select("unit, name")
            if unit:
                query = query.eq("unit", unit)
            res = query.execute()
            for item in res.data:
                item["is_staff"] = False
            return res.data
        except:
            raise HTTPException(status_code=500, detail=str(e))

@app.post("/api/update_personnel_status")
async def update_personnel_status(data: StaffStatusUpdate):
    """Оновлює статус штатності оператора з урахуванням лімітів та видаляє дублікати."""
    clean_name = data.name.strip()
    clean_unit = data.unit.strip()
    
    print(f"DEBUG: START update_personnel_status: Unit='{clean_unit}', Name='{clean_name}', to_staff={data.is_staff}")
    try:
        # 1. Знаходимо ВСІ записи для цього оператора ПЕРЕД змінами
        res_existing = supabase.table("operator_passwords") \
            .select("*") \
            .eq("unit", clean_unit) \
            .eq("name", clean_name) \
            .execute()
        
        existing_records = res_existing.data or []
        print(f"DEBUG: Found {len(existing_records)} existing records for this operator.")
        
        is_already_staff = any(r.get("is_staff") is True for r in existing_records)
        
        # 2. Перевірка ліміту
        unit_lower = clean_unit.lower()
        limit = 3
        if "віпс" in unit_lower:
            limit = 2
        
        if data.is_staff:
            res_all_staff = supabase.table("operator_passwords") \
                .select("name, is_staff") \
                .eq("unit", clean_unit) \
                .eq("is_staff", True) \
                .execute()
            
            # Рахуємо унікальних штатних людей (крім поточного)
            other_staff_names = set(r.get("name") for r in res_all_staff.data if r.get("name") != clean_name)
            current_unique_staff_count = len(other_staff_names)
            
            if not is_already_staff and current_unique_staff_count >= limit:
                unit_type = "ВІПС" if limit == 2 else "ВПС"
                print(f"DEBUG: LIMIT REACHED. Staff already: {list(other_staff_names)}")
                raise HTTPException(status_code=400, detail=f"Ліміт штатних операторів для {unit_type} вичерпано ({limit} чол.)")

        # 3. ВИКОНУЄМО ОНОВЛЕННЯ
        if not existing_records:
            print("DEBUG: Inserting new record (no existing found)")
            ins_res = supabase.table("operator_passwords").insert({
                "unit": clean_unit,
                "name": clean_name,
                "is_staff": data.is_staff,
                "password": "---"
            }).execute()
            print(f"DEBUG: Inserted: {ins_res.data}")
        else:
            # Оновлюємо ВСІ знайдені записи (на випадок дублів)
            upd_res = supabase.table("operator_passwords") \
                .update({"is_staff": data.is_staff}) \
                .eq("unit", clean_unit) \
                .eq("name", clean_name) \
                .execute()
            print(f"DEBUG: Updated matches: {len(upd_res.data)} records.")
            
            # Якщо дублів більше одного — видалимо зайві, щоб навести лад
            if len(existing_records) > 1:
                first_id = existing_records[0]["id"]
                other_ids = [r["id"] for r in existing_records[1:]]
                print(f"DEBUG: Cleaning up {len(other_ids)} duplicates...")
                for oid in other_ids:
                    supabase.table("operator_passwords").delete().eq("id", oid).execute()

        return {"status": "ok", "is_staff": data.is_staff}
    except HTTPException:
        raise
    except Exception as e:
        print(f"CRITICAL ERROR in update_personnel_status: {e}")
        raise HTTPException(status_code=500, detail=f"Помилка оновлення статусу: {str(e)}")

@app.delete("/api/delete_drone/{id}")
async def delete_drone(id: int):
    try:
        supabase.table("drones").delete().eq("id", id).execute()
        return {"status": "ok"}
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

@app.post("/api/add_flight")
async def add_flight(entry: FlightEntry):
    try:
        data = entry.dict()
        # Нормалізація імені оператора перед збереженням
        data["operator"] = normalize_operator_name(data.get("operator", ""))
        
        if entry.result == "Польоти не здійснювались":
            data["duration"] = 0
            data["distance"] = 0
            data["battery_cycles"] = 0
        else:
            data["duration"] = str(calculate_duration(entry.takeoff, entry.landing))
        if "id" in data: del data["id"]
        res = supabase.table("flights").insert(data).execute()
        return {"status": "success", "data": res.data}
    except Exception as e:
        print(f"Database Error: {e}")
        raise HTTPException(status_code=500, detail=str(e))

@app.post("/api/publish_with_telegram")
async def publish_report(report_text: str = Form(...), images: List[UploadFile] = File(None)):
    try:
        async with httpx.AsyncClient() as client:
            if not images:
                await client.post(
                    f"https://api.telegram.org/bot{TELEGRAM_TOKEN}/sendMessage",
                    data={"chat_id": TELEGRAM_CHAT_ID, "text": report_text, "parse_mode": "HTML"}
                )
            else:
                media = []
                files = {}
                for i, img in enumerate(images):
                    file_id = f"pic{i}"
                    img_content = await img.read()
                    files[file_id] = (img.filename, img_content)
                    
                    media_item = {
                        "type": "photo",
                        "media": f"attach://{file_id}"
                    }
                    # Додаємо підпис тільки до першого фото
                    if i == 0:
                        media_item["caption"] = report_text
                        media_item["parse_mode"] = "HTML"
                    media.append(media_item)
                
                # Відправка медіагрупи всередині контекстного менеджера
                await client.post(
                    f"https://api.telegram.org/bot{TELEGRAM_TOKEN}/sendMediaGroup",
                    data={"chat_id": TELEGRAM_CHAT_ID, "media": json.dumps(media)},
                    files=files
                )
        return {"status": "ok"}
    except Exception as e:
        print(f"Telegram API Error: {e}")
        return {"status": "error", "message": str(e)}

@app.get("/api/get_options")
async def get_options():
    now = datetime.now().timestamp()
    if _options_cache["data"] and (now - _options_cache["timestamp"] < _CACHE_TTL):
        return _options_cache["data"]
        
    data = {
        "units": UNITS, 
        "weather": ["Нормальні", "Складні умови", "Несприятливі умови"], 
        "flight_modes": ["Normal", "АТТІ"], 
        "results": ["Без ознак порушення", "Затримання", "Польоти не здійснювались"]
    }
    _options_cache["data"] = data
    _options_cache["timestamp"] = now
    return data

@app.get("/api/get_stats_summary")
async def get_stats_summary(units: Optional[str] = Query(None)):
    """Повертає короткий підсумок для дашбордів без завантаження всіх логів."""
    try:
        query = supabase.table("flights").select("duration, result, unit")
        if units:
            units_list = [u.strip() for u in units.split(",")]
            query = query.in_("unit", units_list)
        
        res = query.execute()
        data = res.data
        
        total_flights = len([f for f in data if f['result'] != "Польоти не здійснювались"])
        total_duration = sum([int(f['duration']) for f in data if f['duration']])
        detentions = len([f for f in data if f['result'] == "Затримання"])
        
        return {
            "total_flights": total_flights,
            "total_duration": total_duration,
            "detentions": detentions
        }
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

@app.get("/api/get_units_drones")
async def get_units_drones(units: str = Query(...)):
    """Отримує дрони для спислу підрозділів (через кому)."""
    units_list = [u.strip() for u in units.split(",")]
    res = supabase.table("drones").select("*").in_("unit", units_list).execute()
    return res.data

@app.get("/api/get_all_flights")
async def get_all_flights(
    unit: Optional[str] = None, 
    units: Optional[str] = Query(None),
    operator: Optional[str] = None,
    from_dt: Optional[str] = None,
    to_dt: Optional[str] = None,
    limit: Optional[int] = Query(None), 
    offset: Optional[int] = Query(0)
):
    query = supabase.table("flights").select("*", count="exact").order("id", desc=True)
    
    if unit:
        query = query.eq("unit", unit)
    elif units:
        units_list = [u.strip() for u in units.split(",")]
        query = query.in_("unit", units_list)
        
    if operator:
        query = query.ilike("operator", f"%{operator}%")
    
    if from_dt:
        # Expected format: "YYYY-MM-DD" or "YYYY-MM-DDTHH:MM"
        query = query.gte("date", from_dt.split('T')[0])
    if to_dt:
        query = query.lte("date", to_dt.split('T')[0])
    
    if limit is not None:
        res = query.range(offset, offset + limit - 1).execute()
        return {
            "data": res.data,
            "total": res.count
        }
    
    # Fallback for old pages
    all_data = []
    batch_limit = 1000
    start = offset
    while True:
        batch_query = supabase.table("flights").select("*").order("id", desc=True).range(start, start + batch_limit - 1)
        if unit:
            batch_query = batch_query.eq("unit", unit)
        elif units:
            units_list = [u.strip() for u in units.split(",")]
            batch_query = batch_query.in_("unit", units_list)
        
        if operator:
            batch_query = batch_query.ilike("operator", f"%{operator}%")
        
        if from_dt:
            batch_query = batch_query.gte("date", from_dt.split('T')[0])
        if to_dt:
            batch_query = batch_query.lte("date", to_dt.split('T')[0])
            
        res = batch_query.execute()
        batch = res.data
        if not batch: break
        all_data.extend(batch)
        if len(batch) < batch_limit: break
        start += batch_limit
        if limit and len(all_data) >= limit: break
    return all_data

@app.delete("/api/delete_flight/{id}")
async def delete_flight(id: int):
    supabase.table("flights").delete().eq("id", id).execute()
    return {"status": "deleted"}

@app.get("/favicon.ico", include_in_schema=False)
async def favicon():
    icon_path = os.path.join(FRONTEND_DIR, "icon.png")
    if os.path.exists(icon_path):
        return FileResponse(icon_path)
    return Response(status_code=204)

# --- PAGE ROUTES ---
@app.get("/")
async def read_index(): return FileResponse(os.path.join(FRONTEND_DIR, "index.html"))
@app.get("/dashboard")
async def read_dashboard(): return FileResponse(os.path.join(FRONTEND_DIR, "dashboard.html"))
@app.get("/request")
async def read_request(): return FileResponse(os.path.join(FRONTEND_DIR, "request.html"))
@app.get("/admin")
async def read_admin(): return FileResponse(os.path.join(FRONTEND_DIR, "admin.html"))
@app.get("/analytics")
async def read_analytics(): return FileResponse(os.path.join(FRONTEND_DIR, "analytics.html"))
@app.get("/report")
async def read_report():
    return FileResponse(
        os.path.join(FRONTEND_DIR, "report.html"),
        headers={
            "Cache-Control": "no-cache, no-store, must-revalidate",
            "Pragma": "no-cache",
            "Expires": "0"
        }
    )
@app.get("/handbook")
async def read_handbook(): return FileResponse(os.path.join(FRONTEND_DIR, "handbook.html"))
@app.get("/fleet")
async def read_fleet(): return FileResponse(os.path.join(FRONTEND_DIR, "fleet_management.html"))
@app.get("/admin_analytics")
async def read_admin_analytics(): return FileResponse(os.path.join(FRONTEND_DIR, "admin_analytics.html"))
@app.get("/support")
async def read_support(): return FileResponse(os.path.join(FRONTEND_DIR, "support.html"))
@app.get("/xxx")
async def read_xxx(): return FileResponse(os.path.join(FRONTEND_DIR, "xxx.html"))

# --- 9. API для ШІ-чату з Інтеграцією Google Maps та Weather API ---
class ChatMessage(BaseModel):
    message: str

@app.post("/api/chat")
async def chat_with_ai(message: str = Form(...), image: Optional[UploadFile] = File(None)):
    user_msg = message.strip()
    
    # 1. Пошук координат у повідомленні (формат 48.4647, 35.0461)
    coords_match = re.search(r"(\d{2}\.\d{3,})[,\s]+(\d{2}\.\d{3,})", user_msg)
    
    context_addon = ""
    if coords_match:
        lat, lon = coords_match.groups()
        google_api_key = os.environ.get("GOOGLE_API_KEY") # ПЕРЕКОНАЙТЕСЬ, ЩО ТУТ НОВИЙ КЛЮЧ
        
        try:
            async with httpx.AsyncClient() as client:
                location_info = "Не визначено"
                msl_info = "Не визначено"
                nearby_places = "Не знайдено"
                
                if google_api_key:
                    # 1. Пряма адреса (Geocoding)
                    geo_url = f"https://maps.googleapis.com/maps/api/geocode/json?latlng={lat},{lon}&key={google_api_key}&language=uk"
                    geo_res = await client.get(geo_url)
                    if geo_res.status_code == 200:
                        geo_data = geo_res.json()
                        if geo_data.get("results"):
                            location_info = geo_data["results"][0].get("formatted_address", "Невідома місцевість")

                    # 2. Висота над рівнем моря MSL (Elevation API)
                    elev_url = f"https://maps.googleapis.com/maps/api/elevation/json?locations={lat},{lon}&key={google_api_key}"
                    elev_res = await client.get(elev_url)
                    if elev_res.status_code == 200:
                        elev_data = elev_res.json()
                        if elev_data.get("results"):
                            msl_info = f"{round(elev_data['results'][0].get('elevation', 0))} м"

                    # 3. Найближчі орієнтири (Places API - прибрано жорсткі фільтри для кращого пошуку)
                    places_url = f"https://maps.googleapis.com/maps/api/place/nearbysearch/json?location={lat},{lon}&radius=5000&key={google_api_key}&language=uk"
                    places_res = await client.get(places_url)
                    if places_res.status_code == 200:
                        places_data = places_res.json()
                        places_list = [p.get("name") for p in places_data.get("results", [])[:5]]
                        if places_list:
                            nearby_places = ", ".join(places_list)

                # 4. Погода (Додано вологість та витягування напрямку вітру)
                weather_url = f"https://api.open-meteo.com/v1/forecast?latitude={lat}&longitude={lon}&current=temperature_2m,relative_humidity_2m,wind_speed_10m,wind_direction_10m,wind_gusts_10m&wind_speed_unit=ms&timezone=auto"
                w_res = await client.get(weather_url)
                weather_info = "Дані недоступні"
                if w_res.status_code == 200:
                    w_data = w_res.json().get("current", {})
                    temp = w_data.get("temperature_2m", "-")
                    hum = w_data.get("relative_humidity_2m", "-")
                    wind = w_data.get("wind_speed_10m", "-")
                    gusts = w_data.get("wind_gusts_10m", "-")
                    wind_dir = w_data.get("wind_direction_10m", "-")
                    weather_info = f"Температура: {temp}°C, Вологість: {hum}%, Вітер: {wind} м/с (пориви {gusts} м/с), Напрямок вітру: {wind_dir}°"

                # 5. Магнітні бурі (K-index) з офіційного API NOAA (США)
                k_index = "Невідомо"
                try:
                    noaa_url = "https://services.swpc.noaa.gov/products/noaa-planetary-k-index.json"
                    noaa_res = await client.get(noaa_url, timeout=3.0)
                    if noaa_res.status_code == 200:
                        noaa_data = noaa_res.json()
                        if len(noaa_data) > 1:
                            k_index = noaa_data[-1][1] # Беремо найсвіжіший Kp індекс
                except Exception as e:
                    print(f"NOAA K-index Error: {e}")

                now_date = datetime.now().strftime("%d.%m.%Y")
                now_time = datetime.now().strftime("%H:%M")
                
                # Додаток до промпту: ЖОРСТКІ ІНСТРУКЦІЇ
                context_addon = f"""
                
                [СИСТЕМА: ВИЯВЛЕНО КООРДИНАТИ {lat}, {lon}. НАДАЮ АВТОМАТИЧНІ ДАНІ]
                - Топографія: Адреса: {location_info}. Висота MSL: {msl_info}. Орієнтири (5км): {nearby_places}.
                - Метео: {weather_info}. K-index: {k_index}.
                
                ІНСТРУКЦІЯ ДЛЯ ШІ ЩОДО ФОРМАТУВАННЯ ЗВІТУ:
                1. РЕЛЬЄФ: Якщо рельєф не вказано прямо, проаналізуй висоту {msl_info} та адресу і зроби логічний висновок (напр. "Рівнина", "Височина", "Міська забудова"). НЕ ПИШИ "Дані відсутні".
                2. НАПРЯМОК ВІТРУ: Конвертуй азимут вітру в сторони світу (Пн, Пд, Зх, Сх) і вкажи, куди буде зносити дрон (ЗНЕСЕННЯ ATTI).
                3. Заповни всі дані без винятків.
                """
        except Exception as e:
            print(f"API Fetch Error: {e}")

    # Збираємо фінальний текст для ШІ
    final_prompt = user_msg + context_addon

    system_prompt = os.environ.get("AI_SYSTEM_PROMPT", "Ти — терміновий технічний асистент та інструктор із БпЛА (DJI, Autel), РЕБ/РЕР. Твій пріоритет — миттєва допомога оператору. ПРАВИЛА: 1. Пиши коротко, без довгих вступів. 2. Надавай чіткий алгоритм дій. 3. Аналізуй локації та погоду для планування маршрутів, якщо користувач надсилає координати.")
    
    async def generate_response():
        try:
            if not ai_client:
                yield "Дякую за запитання! Будь ласка, налаштуйте API-ключ Gemini у файлі .env."
                return

            contents = []
            
            # 1. Додаємо всі PDF-мануали з бази знань
            if knowledge_files_cache:
                contents.extend(knowledge_files_cache)
                
            # 2. Додаємо сам запит користувача з координатами
            contents.append(final_prompt)
            if image:
                image_bytes = await image.read()
                contents.append(types.Part.from_bytes(data=image_bytes, mime_type=image.content_type))
            
            model_name = os.environ.get("GEMINI_MODEL_NAME", "gemini-flash-latest")
            
            # Вмикаємо функцію Grounding (пошук в Google), як ви і вказали
            ai_config = types.GenerateContentConfig(
                system_instruction=system_prompt,
                tools=[{"google_search": {}}]
            )
            
            try:
                response = await ai_client.aio.models.generate_content_stream(
                    model=model_name,
                    contents=contents,
                    config=ai_config
                )
                async for chunk in response:
                    if chunk.text:
                        yield chunk.text
            except Exception as e:
                err_str = str(e)
                print(f"Primary Model Error: {err_str}")
                # If the error is about documents with no pages, retry without knowledge files
                if "no pages" in err_str.lower() or "не має сторінок" in err_str.lower() or "document" in err_str.lower():
                    print("⚠️ Помилка бази знань — спроба без документів...")
                    safe_contents = [c for c in contents if not hasattr(c, 'uri')]
                else:
                    safe_contents = contents
                # Fallback without grounding, using safe_contents (no bad docs)
                fallback_config = types.GenerateContentConfig(system_instruction=system_prompt)
                response = await ai_client.aio.models.generate_content_stream(
                    model="gemini-2.5-flash-lite",
                    contents=safe_contents,
                    config=fallback_config
                )
                async for chunk in response:
                    if chunk.text:
                        yield chunk.text
        except Exception as e:
            print(f"AI Stream Error: {e}")
            yield "Сервіс ШІ тимчасово недоступний (високе навантаження або вичерпано ліміти)."

    return StreamingResponse(generate_response(), media_type="text/plain")

class DocxRequest(BaseModel):
    text: str
    filename: str

# In-memory cache for generated DOCX files (cleared after download or timeout)
_docx_cache: dict = {}

@app.post("/api/generate_docx")
async def generate_docx(report_data: str = Form(...), filename: str = Form(...)):
    print(f"Generating DOCX: {filename}")
    try:
        from docx.shared import Cm, Pt
        try:
            data = json.loads(report_data)
        except json.JSONDecodeError as je:
            print(f"JSON Decode Error: {je}")
            raise HTTPException(status_code=400, detail=f"Invalid JSON data: {str(je)}")

        print(f"Report Data Keys: {list(data.keys())}")
        if 'flights' in data:
            print(f"Flights count: {len(data['flights'])}")
            
        # ──────────────────────────────────────────────────────────────────────────
    # ЗАМІНИТИ весь блок generate_docx починаючи від "doc = Document()" і до
    # кінця функції (перед "except Exception as e:") на код нижче.
    # ──────────────────────────────────────────────────────────────────────────

        doc = Document()

        # --- НАЛАШТУВАННЯ ПОЛІВ ---
        section = doc.sections[0]
        section.top_margin = Cm(1.5)      # Верхнє поле
        section.bottom_margin = Cm(1.0)   # Нижнє поле
        section.left_margin = Cm(2.0)     # Ліве поле (зазвичай більше для підшивання)
        section.right_margin = Cm(2.0)    # Праве поле
        # -------------------------

        # --- Налаштування шрифту (Times New Roman, 12pt) ---
        style = doc.styles['Normal']
        font = style.font
        font.name = 'Times New Roman'
        font.size = Pt(12)
        
        # Видаляємо інтервали між абзацами та встановлюємо одинарний інтервал
        from docx.enum.text import WD_LINE_SPACING, WD_TAB_LEADER, WD_TAB_ALIGNMENT, WD_ALIGN_PARAGRAPH
        style.paragraph_format.space_after = Pt(0)
        style.paragraph_format.space_before = Pt(0)
        style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
        r = style.element.rPr.rFonts
        r.set(qn('w:ascii'), 'Times New Roman')
        r.set(qn('w:hAnsi'), 'Times New Roman')

        # --- Хелпер: абзац з жирним підкресленим лейблом + звичайний текст ---
        def add_labeled(label: str, value: str = "", align=None, newline_before=False):
            """Додає абзац: [жирний підкреслений label][звичайний value]."""
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.space_before = Pt(0)
            if align:
                p.alignment = align
            if newline_before:
                p.add_run("\n")
            run_label = p.add_run(label)
            run_label.bold = True
            run_label.underline = True
            if value:
                p.add_run(value)
            return p

        # --- Хелпер: лише звичайний абзац ---
        def add_plain(text: str, align=None):
            p = doc.add_paragraph(text)
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.space_before = Pt(0)
            if align:
                p.alignment = align
            return p

        # 1. Шапка документа (вирівняна праворуч)
        from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
        header_p = doc.add_paragraph()

        # Налаштовуємо форматування параграфа
        fmt = header_p.paragraph_format
        # 1. Робимо відступ зліва 9 см (як на лінійці в скріншоті)
        fmt.left_indent = Cm(9.5)
        # 2. Додаємо точку табуляції на правому краї (наприклад, 16.5 см)
        # Це дозволить рознести "підполковнику" та "Армену МКРТЧЯН"
        fmt.tab_stops.add_tab_stop(Cm(17.5), WD_TAB_ALIGNMENT.RIGHT)

        # Важливо: прибираємо WD_ALIGN_PARAGRAPH.RIGHT, 
        # щоб текст всередині блоку був вирівняний по лівому краю
        header_p.alignment = WD_ALIGN_PARAGRAPH.LEFT 

        # Формуємо текст
        # Зверніть увагу на символ \t (табуляція) в останньому рядку
        header_text = data.get(
            'header',
            "Начальнику відділу організації повітряної\n"
            "розвідки та протидії безпілотним\n"
            "повітряним суднам штабу\n"
            "підполковнику\tАрмену МКРТЧЯН"
        )

        header_p.add_run(header_text)

        # 2. Назва документа (по центру, жирна)
        doc.add_paragraph() # Порожній рядок перед заголовком
        title_p = doc.add_paragraph()
        title_p.paragraph_format.space_after = Pt(0)
        title_p.paragraph_format.space_before = Pt(0)
        title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run = title_p.add_run("ДОНЕСЕННЯ ПРО ПОЛІТ")
        title_run.bold = True
        doc.add_paragraph() # Порожній рядок після заголовка

        # 3. Базові дані — лейбли жирні+підкреслені
        add_labeled("Дата вильотів", f": {data.get('date', '__.__.____')}")
        vips = data.get('vips', '')
        vps = data.get('sector_vps', '')

        sector_parts = []
        if vips:
            sector_parts.append(f"віпс {vips}")
        if vps:
            sector_parts.append(f"впс {vps}")
        sector_display = ' '.join(sector_parts) if sector_parts else '___'

        add_labeled("Ділянка", f": {sector_display}")
        add_labeled("Район виконання завдань", f": {data.get('task_area_coords', '___')}")
        add_labeled("Вильоти", ":")

        # 4. Таблиця польотів
        flights = data.get('flights', [])
        table = doc.add_table(rows=1, cols=4)
        table.style = 'Table Grid'
        # Налаштування ширини буде виконано нижче, після додавання всіх рядків
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Прізвище екіпажу'
        hdr_cells[1].text = 'К-сть'
        hdr_cells[2].text = 'Тип БпАК'
        hdr_cells[3].text = 'Час та дистанція вильотів'

        if not flights:
            row_cells = table.add_row().cells
            row_cells[0].text = 'Немає даних'
            row_cells[1].text = '0'
            row_cells[2].text = '-'
            row_cells[3].text = '-'
        else:
            for flight in flights:
                row_cells = table.add_row().cells
                row_cells[0].text = flight.get('operator', '')
                row_cells[1].text = str(flight.get('count', ''))
                row_cells[2].text = flight.get('drone', '')
                import re as _re
                details_raw = flight.get('details', '')
                details_clean = _re.sub(r'(\d{2}:\d{2}):\d{2}', r'\1', details_raw)
                # Split sorties and put each on a new line inside the cell
                sorties = [s.strip() for s in details_clean.split(';') if s.strip()]
                cell = row_cells[3]
                cell.text = ''  # clear default paragraph
                for i, sortie in enumerate(sorties):
                    if i == 0:
                        cell.paragraphs[0].text = sortie
                    else:
                        cell.add_paragraph(sortie)
        # Фіксуємо ширину таблиці та колонок ПІСЛЯ додавання всіх рядків:
        table.autofit = False 
        table.width = Cm(17)
        widths = [Cm(4.0), Cm(2.0), Cm(6.0), Cm(5.0)]
        for i, width in enumerate(widths):
            table.columns[i].width = width
        for row in table.rows:
            for i, width in enumerate(widths):
                row.cells[i].width = width

        # 5. Маршрут та екіпаж
        add_labeled("Маршрут", f": {data.get('route', '___')}")
        add_labeled("БпАК", f": {data.get('drones_list', '___')}")
        add_labeled("Склад екіпажу", ":")
        add_plain(f"Оператор: {data.get('commander', '___')};")

        # 6. Результати
        add_labeled("Результати", ":")
        add_plain(data.get('result', "Під час польотів порушень ОПДК не виявлено."))

        # 7. Фотофіксація (якщо є фото)
        photo_b64 = data.get('photo')
        temp_photo_path = None
        if photo_b64 and photo_b64.startswith('data:image'):
            try:
                import base64
                header_data, encoded = photo_b64.split(",", 1)
                photo_data = base64.b64decode(encoded)
                temp_photo_path = os.path.join("app", f"{uuid.uuid4()}.jpg")
                with open(temp_photo_path, "wb") as f:
                    f.write(photo_data)
                pic_para = doc.add_paragraph()
                pic_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                pic_para.add_run().add_picture(temp_photo_path, width=Cm(11))
            except Exception as e:
                print(f"Помилка завантаження фото: {e}")

        # 8. Зв'язок (використовуємо той самий sector_display, що й у рядку "Ділянка")
        add_labeled("Зв'язок на кордоні підтримувався", ":")
        comm_sector = f"з ПН, ОЧ {sector_display} " if sector_display != '___' else ''
        add_plain(f"{comm_sector}по р/ст. по радіомережі; відео-, фотодокументування здійснювалися штатною апаратурою БпАК.")

        # 9. Погода та стан техніки
        add_labeled("Погода по маршруту", f": {data.get('weather', '___')}")
        add_labeled("Готовність екіпажу та стан БпАК", ":")
        add_plain("БпАК справний, недоліків у роботі техніки не виявлено. Політ виконано в штатному режимі, відмов у системах керування та телеметрії не зафіксовано. Зауважень немає.")

        # 10. Підпис
        add_labeled("Донесення склав", "")
        add_plain("Оператор:")
        
        # Формуємо рядок підпису без крапок (тільки відступ)
        sig_p = doc.add_paragraph()
        sig_p.paragraph_format.space_after = Pt(0)
        sig_p.paragraph_format.space_before = Pt(0)
        # Додаємо табуляцію БЕЗ лідера
        sig_p.paragraph_format.tab_stops.add_tab_stop(Cm(16.5), WD_TAB_ALIGNMENT.RIGHT)
        
        full_text = data.get('commander_short', '___')
        if "  " in full_text:
            parts = [p.strip() for p in full_text.split("  ") if p.strip()]
            if len(parts) >= 2:
                rank_part = parts[0]
                name_part = parts[-1]
                sig_p.add_run(f"{rank_part}\t{name_part}")
            else:
                sig_p.add_run(full_text.replace("  ", "\t"))
        else:
            sig_p.add_run(full_text.replace(" ", "\t", 1))

        # ── Збереження та відповідь ────────────────────────────────────────────
        from io import BytesIO
        from urllib.parse import quote

        file_stream = BytesIO()
        doc.save(file_stream)
        file_bytes = file_stream.getvalue()

        if not filename.lower().endswith(".docx"):
            filename += ".docx"

        encoded_filename = quote(filename)
        safe_ascii = "report.docx"

        return Response(
            content=file_bytes,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={
                "Content-Disposition": (
                    f"attachment; filename=\"{safe_ascii}\"; "
                    f"filename*=UTF-8''{encoded_filename}"
                )
            }
        )

    except Exception as e:
        print(f"DOCX Generation Error: {e}")
        raise HTTPException(status_code=500, detail=str(e))

@app.get("/api/download_docx/{token}/{filename}")
async def download_docx(token: str, filename: str):
    entry = _docx_cache.get(token)
    if not entry:
        raise HTTPException(status_code=404, detail="File not found or expired")
    data = entry['data']
    # Remove from cache after serving
    del _docx_cache[token]
    from io import BytesIO
    from urllib.parse import quote
    
    encoded_filename = quote(filename)
    safe_ascii = "report.docx"
    
    return StreamingResponse(
        BytesIO(data),
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f"attachment; filename=\"{safe_ascii}\"; filename*=UTF-8''{encoded_filename}"}
    )

@app.post("/api/weather_analysis")
async def weather_analysis(req: WeatherAnalysisRequest):
    try:
        url = "https://api.open-meteo.com/v1/forecast"
        params = {
            "latitude": req.lat,
            "longitude": req.lon,
            "hourly": ["temperature_2m", "apparent_temperature", "precipitation_probability", "precipitation", "visibility", "wind_speed_80m", "wind_speed_120m", "wind_speed_180m", "wind_direction_80m", "wind_direction_120m", "wind_direction_180m", "wind_gusts_10m", "temperature_80m", "temperature_120m", "temperature_180m", "wind_speed_10m", "pressure_msl", "dew_point_2m", "relative_humidity_2m", "wind_direction_10m", "cloud_cover"],
            "wind_speed_unit": "ms",
            "start_date": req.date,
            "end_date": req.date,
            "timezone": "auto"
        }
        
        responses = await asyncio.to_thread(openmeteo.weather_api, url, params=params)
        response = responses[0]
        
        hourly = response.Hourly()
        hourly_data = {"date": pd.date_range(
            start = pd.to_datetime(hourly.Time(), unit = "s", utc = True),
            end =  pd.to_datetime(hourly.TimeEnd(), unit = "s", utc = True),
            freq = pd.Timedelta(seconds = hourly.Interval()),
            inclusive = "left"
        )}

        hourly_data["temperature_2m"] = hourly.Variables(0).ValuesAsNumpy()
        hourly_data["apparent_temperature"] = hourly.Variables(1).ValuesAsNumpy()
        hourly_data["precipitation_probability"] = hourly.Variables(2).ValuesAsNumpy()
        hourly_data["precipitation"] = hourly.Variables(3).ValuesAsNumpy()
        hourly_data["visibility"] = hourly.Variables(4).ValuesAsNumpy()
        hourly_data["wind_speed_80m"] = hourly.Variables(5).ValuesAsNumpy()
        hourly_data["wind_speed_120m"] = hourly.Variables(6).ValuesAsNumpy()
        hourly_data["wind_speed_180m"] = hourly.Variables(7).ValuesAsNumpy()
        hourly_data["wind_direction_80m"] = hourly.Variables(8).ValuesAsNumpy()
        hourly_data["wind_direction_120m"] = hourly.Variables(9).ValuesAsNumpy()
        hourly_data["wind_direction_180m"] = hourly.Variables(10).ValuesAsNumpy()
        hourly_data["wind_gusts_10m"] = hourly.Variables(11).ValuesAsNumpy()
        hourly_data["temperature_80m"] = hourly.Variables(12).ValuesAsNumpy()
        hourly_data["temperature_120m"] = hourly.Variables(13).ValuesAsNumpy()
        hourly_data["temperature_180m"] = hourly.Variables(14).ValuesAsNumpy()
        hourly_data["wind_speed_10m"] = hourly.Variables(15).ValuesAsNumpy()
        hourly_data["pressure_msl"] = hourly.Variables(16).ValuesAsNumpy()
        hourly_data["dew_point_2m"] = hourly.Variables(17).ValuesAsNumpy()
        hourly_data["relative_humidity_2m"] = hourly.Variables(18).ValuesAsNumpy()
        hourly_data["wind_direction_10m"] = hourly.Variables(19).ValuesAsNumpy()
        hourly_data["cloud_cover"] = hourly.Variables(20).ValuesAsNumpy()

        hourly_dataframe = pd.DataFrame(data = hourly_data)
        
        # Adjust UTC to Local Time to match req.time
        utc_offset_seconds = response.UtcOffsetSeconds()
        hourly_dataframe['local_time'] = hourly_dataframe['date'] + pd.Timedelta(seconds=utc_offset_seconds)
        
        req_hour = int(req.time.split(":")[0])
        target_row = hourly_dataframe[hourly_dataframe['local_time'].dt.hour == req_hour]
        
        if not target_row.empty:
            weather_data = target_row.iloc[0].to_dict()
            weather_data['date'] = str(weather_data['date'])
            weather_data['local_time'] = str(weather_data['local_time'])
        else:
            weather_data = hourly_dataframe.iloc[0].to_dict() # Fallback to first hour
            weather_data['date'] = str(weather_data['date'])
            weather_data['local_time'] = str(weather_data['local_time'])

        # UAV specs for AI context
        uav_specs = {
            "DJI Mavic 3":    "Макс. вітростійкість: 12 м/с. Захист: IP43. Мін. темп.: -10°C.",
            "DJI Matrice 4T": "Макс. вітростійкість: 15 м/с. Захист: IP55. Мін. темп.: -20°C.",
            "DJI Matrice 30T": "Макс. вітростійкість: 15 м/с. Захист: IP55. Мін. темп.: -20°C.",
            "DJI Matrice 300": "Макс. вітростійкість: 15 м/с. Захист: IP45. Мін. темп.: -20°C.",
            "Autel EVO Max 4T": "Макс. вітростійкість: 15 м/с. Захист: IP43. Мін. темп.: -20°C.",
        }
        uav_spec_info = uav_specs.get(req.uav_model, f"Модель: {req.uav_model} (ТТХ не відомі)")

        system_prompt = (
            "Ти — експерт з авіаційної метеорології для БпЛА. "
            "Аналізуй отримані погодні дані та надавай структуровану відповідь у форматі JSON. "
            "Враховуй ТТХ зазначеного БпЛА (макс. вітростійкість, мін. температура, пориви вітру) при оцінці ризиків. "
            "Додай оцінку (ІДЕАЛЬНО/НОРМА/НЕСПРИЯТЛИВО/НЕПРИЙНЯТНО) до поля вітер і поривів вітру, зіставляючи з макс.свого БпЛА."
        )
        
        user_prompt = f"""
Аналізуй погодні умови для польоту БпЛА:
Координати: {req.lat}, {req.lon}
Час польоту: {req.date} {req.time}
Модель БпЛА: {req.uav_model}
ТТХ БпЛА: {uav_spec_info}
Висота польоту: {req.altitude} м

Дані погодної моделі (Open-Meteo SDK):
{json.dumps(weather_data, indent=2, ensure_ascii=False)}

ЗАВДАННЯ:
1. `wind_speed_10m` + `wind_gusts_10m` + `wind_direction_10m` — вітер біля землі (додай оцінку ІДЕАЛЬНО/НОРМА/НЕСПРИЯТЛИВО відносно ТТХ БпЛА).
2. `wind_gusts_10m` — окремий пункт wind_gusts (оціни безпеку зіставляючи з макс.вітростійкістью БпЛА).
3. `wind_speed_80m/120m/180m` і `wind_direction_80m/120m/180m` — вітер по висоті.
4. `temperature_80m/120m/180m` — ризик обледеніння з урах.мін.темп.експлуатації БпЛА.
5. `cloud_cover` — хмарність у відсотках.
6. `pressure_msl` — щільнісна висота.
7. Всі числові дані конвертуй у зрозумілий формат (напр. 5.2 м/с).

ПОВЕРНИ ВІДПОВІДЬ СУВОРО В JSON. НЕ ЗМІНЮЙ КЛЮЧІ:
{{
  "wind_surface": "швидкість, напрямок, оцінка (ІДЕАЛЬНО/НОРМА/НЕСПРИЯТЛИВО)",
  "wind_gusts": "пориви м/с, оцінка безпеки для цього БпЛА",
  "wind_altitude": {{
    "80m": "швидкість і напрямок",
    "120m": "швидкість і напрямок",
    "180m": "швидкість і напрямок",
    "target": "прогноз для {req.altitude}м"
  }},
  "cloudiness": "хмарність % (cloud_cover)",
  "visibility": "видимість в м або км",
  "precipitation": "опади та ймовірність",
  "density_altitude": "щільнісна висота в м",
  "icing_risk": "ризик (Низький/Середній/Високий)",
  "wind_shear": "зсув вітру (різниця між висотами)",
  "pessimistic_estimate": "песимістична оцінка",
  "confidence": {{
    "score": 92,
    "summary": "короткий опис",
    "details": {{
      "fog": 95, "precipitation": 92, "temperature": 94, "visibility": 88, "wind": 90
    }}
  }}
}}
"""
        if not ai_client:
            raise HTTPException(status_code=500, detail="AI Client not configured")

        response_ai = await ai_client.aio.models.generate_content(
            model="gemini-2.5-flash-lite",
            contents=user_prompt,
            config=types.GenerateContentConfig(
                system_instruction=system_prompt,
                response_mime_type="application/json"
            )
        )
        
        return json.loads(response_ai.text)

    except Exception as e:
        print(f"Weather Analysis Error: {e}")
        import traceback
        traceback.print_exc()
        raise HTTPException(status_code=500, detail=str(e))

# Serving all static files from root for PWA compatibility
app.mount("/", StaticFiles(directory=FRONTEND_DIR, html=True), name="static")

if __name__ == "__main__":
    import uvicorn
    uvicorn.run(app, host="127.0.0.1", port=8000)

